"""
Build a reformatted, readable Word doc from the Qualtrics export.
Matrix questions are condensed. Proposed removals are highlighted.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, re, html

# ── helpers ─────────────────────────────────────────────────────────────────

YELLOW = RGBColor(0xFF, 0xFF, 0x00)
RED    = RGBColor(0xFF, 0x00, 0x00)
GRAY   = RGBColor(0x60, 0x60, 0x60)

def set_highlight(run, color_hex="FFFF00"):
    """Add word highlight (shading) to a run."""
    rPr = run._r.get_or_add_rPr()
    highlight = OxmlElement('w:highlight')
    highlight.set(qn('w:val'), 'yellow')
    rPr.append(highlight)

def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    return p

def add_para(doc, text="", bold=False, italic=False, size=11,
             color=None, highlight=False, left_indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(left_indent)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        if highlight:
            set_highlight(run)
    return p

def add_bullet(doc, text, level=0, highlight=False, strikethrough=False):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Inches(0.25 * (level + 1))
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(10)
    if highlight:
        set_highlight(run)
    if strikethrough:
        run.font.strike = True
        run.font.color.rgb = GRAY
    return p

def add_scale_note(doc, text, indent=0.5):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(indent)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    return p

def add_qvar(doc, varname, question_text, highlight_block=False):
    """Print variable name in gray + question text bold on same paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)

    r_var = p.add_run(f"[{varname}]  ")
    r_var.font.size = Pt(9)
    r_var.font.color.rgb = GRAY

    r_q = p.add_run(question_text)
    r_q.bold = True
    r_q.font.size = Pt(11)
    if highlight_block:
        set_highlight(r_q)
    return p

def section_break(doc, title):
    doc.add_paragraph()
    p = doc.add_heading(title, level=2)
    p.paragraph_format.space_before = Pt(14)
    return p

def removal_banner(doc, label):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(f"  ★ PROPOSED REMOVAL: {label}  ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    set_highlight(run)
    return p

# ── items proposed for removal ───────────────────────────────────────────────

# research_5 and research_6: full blocks removed
# govtfund_4: keep items 2,3,6,7,8,10 → cut 1,4,5,9,11,12,13
G4_CUT  = {1, 4, 5, 9, 11, 12, 13}
G4_KEEP = {2, 3, 6, 7, 8, 10}

# govtfund_5: keep items 3,4,5,6,7 → cut 1,2,8,9,10
G5_CUT  = {1, 2, 8, 9, 10}
G5_KEEP = {3, 4, 5, 6, 7}

# ── build document ──────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

# Title
t = doc.add_heading("UnivNews Step 1 — Survey Questions", level=1)
t.paragraph_format.space_after = Pt(2)

add_para(doc, "Cohorts 1 & 2  |  Proposed removals highlighted in yellow", italic=True, size=10, color=GRAY)
add_para(doc, "")

# Legend
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(10)
r1 = p.add_run("Highlighting key:  ")
r1.font.size = Pt(10); r1.bold = True
r2 = p.add_run(" ENTIRE BLOCK REMOVED ")
r2.font.size = Pt(10); r2.bold = True
set_highlight(r2)
r3 = p.add_run("  = full block cut.  ")
r3.font.size = Pt(10)
r4 = p.add_run(" individual item ")
r4.font.size = Pt(10)
set_highlight(r4)
r5 = p.add_run("  = specific item cut from that block.")
r5.font.size = Pt(10)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
#  BLOCK 1: CONSENT
# ════════════════════════════════════════════════════════════════════════════
section_break(doc, "Block: Consent")

add_qvar(doc, "consent", "If you consent to participate, please select the first option to begin the study.")
add_bullet(doc, "I agree to participate in this study.")
add_bullet(doc, "I do not agree to participate in this study.")

doc.add_paragraph()
add_qvar(doc, "consent_2", "IMPORTANT: Please complete the survey in one sitting. [Full text re: payment rejection for inattentive responding.]")
add_bullet(doc, "I have read the information above and I understand that my submission may be rejected and I will not receive payment if I do not complete the survey carefully.")

# ════════════════════════════════════════════════════════════════════════════
#  BLOCK 2: RESEARCH
# ════════════════════════════════════════════════════════════════════════════
section_break(doc, "Block: Research")

add_para(doc, "Intro text: The purpose of this study is to understand people's opinions about research conducted at universities. [Continues on screen.]", italic=True, size=9, color=GRAY)

doc.add_paragraph()
add_qvar(doc, "research_1", "How much do you, personally, know about research at colleges and universities?")
add_bullet(doc, "None at all")
add_bullet(doc, "A little")
add_bullet(doc, "A moderate amount")
add_bullet(doc, "A lot")
add_bullet(doc, "A great deal")

doc.add_paragraph()
add_qvar(doc, "research_2", 'To what extent do you agree or disagree with the following statement? "Scientific research methods can yield valuable insights about the world."')
add_bullet(doc, "Strongly disagree → Strongly agree")
add_scale_note(doc, "5-point scale: Strongly disagree / Somewhat disagree / Neither agree nor disagree / Somewhat agree / Strongly agree")

doc.add_paragraph()
add_qvar(doc, "research_3", "To what extent has research at colleges and universities had a positive or negative effect on:")
for item in [
    "You personally", "People you care about", "People in your local community",
    "People in the United States", "Students at universities", "Scientists at universities",
    "Universities themselves", "Pharmaceutical companies", "Technology companies",
    "The military", "The US government", "Wealthy individuals", "Disadvantaged individuals"
]:
    add_bullet(doc, item)
add_scale_note(doc, "5-point scale: Extremely negative / Somewhat negative / Neither positive nor negative / Somewhat positive / Extremely positive")

doc.add_paragraph()
add_qvar(doc, "research_3b", "Please describe the positive or negative effects that academic research has had on you personally. [Open text, minimum 3 sentences]")

doc.add_paragraph()
add_qvar(doc, "research_4", "Rate your agreement with the following statements:")
for item in [
    "I usually understand academic research findings discussed online or in the news",
    "I know where to look up information about academic research",
    "It is clear how to apply findings from academic research in my own life",
    "I often think about the impact of academic science on my everyday life",
    "The benefits of scientific research outweigh harmful results",
    "Science makes our way of life change too fast",
    "Because of scientific research, there will be more opportunities for the next generation",
    "Even if it brings no immediate benefits, scientific research that advances the frontiers of knowledge is necessary",
]:
    add_bullet(doc, item)
add_scale_note(doc, "5-point scale: Strongly disagree / Somewhat disagree / Neither agree nor disagree / Somewhat agree / Strongly agree")

# ── research_5 (FULL REMOVAL) ─────────────────────────────────────────────
doc.add_paragraph()
removal_banner(doc, "entire research_5 block — cut from Cohort 3 onward")
add_qvar(doc, "research_5",
         "Below are some types of scientific research areas and outcomes. For each, please indicate whether that type of research is important for the public good, and whether academic research is likely to produce that area/outcome.",
         highlight_block=True)
for item in [
    "Advancing social science (insights into human behavior, mental health, poverty, misinformation)",
    "Developing new medicines to prevent or cure diseases (cancer, heart disease, dementia)",
    "Identifying the best ways to protect the environment and ensure clean air and water",
    "Developing new energy technologies to make America more energy independent",
    "Developing advanced technologies to drive economic growth and increase wages",
    "Developing innovative manufacturing processes to boost domestic production",
    "Developing new defense capabilities and national security technologies",
    "Advancing artificial intelligence and machine learning",
    "Exploring space to understand the solar system and broader universe",
    "Improvement in education and learning systems and structures",
    "Development of better public policies",
    "Better understanding for individuals about how to improve their lives",
    "More awareness of others' experiences and perspectives",
]:
    add_bullet(doc, item, highlight=True)
add_scale_note(doc,
    "Sub-scale A (importance): 5-point scale: Not at all important / Slightly important / Moderately important / Very important / Extremely important\n"
    "       Sub-scale B (likelihood): 5-point scale: Very unlikely / Somewhat unlikely / Neutral / Somewhat likely / Very likely")

# ── research_6 (FULL REMOVAL) ─────────────────────────────────────────────
doc.add_paragraph()
removal_banner(doc, "entire research_6 block — cut from Cohort 3 onward")
add_qvar(doc, "research_6",
         "Below are potential negative outcomes from academic research. For each, please indicate how problematic that outcome would be, and whether academic research is likely to produce that outcome.",
         highlight_block=True)
for item in [
    "Researchers make mistakes and results are incorrect",
    "Research is slow, and it takes a long time to have an impact",
    "Few people are able to access or understand the results of research",
    "Findings are specific to certain groups of individuals and don't apply to others",
    "Studies fail to yield impactful results",
    "Research is conducted inefficiently and wastes money",
    "New technology resulting from research harms the general public by eliminating jobs",
    "New technology resulting from research invades our privacy",
    "Research can create new technology that causes widespread damage (such as nuclear weapons)",
    "Research findings are misused",
    "The way that results are reported is biased by news outlets",
]:
    add_bullet(doc, item, highlight=True)
add_scale_note(doc,
    "Sub-scale A (problematic): 5-point scale: Not at all problematic / Slightly problematic / Moderately problematic / Very problematic / Extremely problematic\n"
    "       Sub-scale B (likelihood): 5-point scale: Very unlikely / Somewhat unlikely / Neutral / Somewhat likely / Very likely")

# ════════════════════════════════════════════════════════════════════════════
#  BLOCK 3: GOVERNMENT FUNDING
# ════════════════════════════════════════════════════════════════════════════
section_break(doc, "Block: Government Funding")

add_para(doc, "Intro text: As you may know, the federal government invests public funds (money from taxpayers) into conducting research at colleges and universities. [Continues on screen.]", italic=True, size=9, color=GRAY)

doc.add_paragraph()
add_qvar(doc, "govtfund_1", "Do you approve or disapprove of the federal government using taxpayer funds to invest in research at colleges and universities?")
for opt in ["Strongly disapprove", "Somewhat disapprove", "Somewhat approve", "Strongly approve", "Don't know"]:
    add_bullet(doc, opt)

doc.add_paragraph()
add_qvar(doc, "govtfund_1b", "Please describe your viewpoint on federal funding of academic research — why do you approve or disapprove? [Open text, minimum 3 sentences]")

doc.add_paragraph()
add_qvar(doc, "govtfund_2", "Which statement comes closer in your view, even if neither is exactly right?")
for opt in [
    "Government investment in research is ESSENTIAL for scientific progress",
    "Private investment will ensure that enough progress is made, even without government investment",
    "Don't know"
]:
    add_bullet(doc, opt)

doc.add_paragraph()
add_qvar(doc, "govtfund_3", "How much do you currently know about cuts to government funding for academic research?")
for opt in ["None at all", "A little", "A moderate amount", "A lot", "A great deal"]:
    add_bullet(doc, opt)

# ── govtfund_4 (PARTIAL REMOVAL) ─────────────────────────────────────────
doc.add_paragraph()
g4_items = [
    (1,  "People who will directly benefit from the research"),
    (2,  "Students at universities"),
    (3,  "Scientists at universities"),
    (4,  "The general public"),
    (5,  "People in the military"),
    (6,  "People who identify as liberals"),
    (7,  "People who identify as conservatives"),
    (8,  "People who favor minimal government spending"),
    (9,  "Politicians"),
    (10, "People who don't believe in mainstream science"),
    (11, "Medical and mental health professionals"),
    (12, "Pharmaceutical companies"),
    (13, "Technology companies"),
]
add_para(doc, f"  ★ PROPOSED PARTIAL REDUCTION: govtfund_4 — keep 6 of 13 items (items 2, 3, 6, 7, 8, 10 in bold; others highlighted for removal)",
         bold=True, size=10, color=RGBColor(0xCC, 0x00, 0x00))
add_qvar(doc, "govtfund_4",
         "How do you think each of these groups of people feel about the government financially supporting academic research?")
for num, text in g4_items:
    cut = num in G4_CUT
    add_bullet(doc, text, highlight=cut, strikethrough=cut)
add_scale_note(doc, "5-point scale: Strongly disapprove / Somewhat disapprove / Neither approve nor disapprove / Somewhat approve / Strongly approve")

# ── govtfund_5 (PARTIAL REMOVAL) ─────────────────────────────────────────
doc.add_paragraph()
g5_items = [
    (1,  "Researchers have enough resources to explore their ideas"),
    (2,  "Research can focus on long-term projects instead of focusing on making an immediate profit"),
    (3,  "Exciting research is sped up"),
    (4,  "Results will become publicly available"),
    (5,  "Government-funded researchers are held to higher standards than privately funded work"),
    (6,  "Money is spent on research that should have been spent on more important things"),
    (7,  "Taxpayer money is wasted"),
    (8,  "Research is delayed by government oversight"),
    (9,  "Political agendas influence what research gets funded"),
    (10, "Researchers lie about their results to get continued funding"),
]
add_para(doc, f"  ★ PROPOSED PARTIAL REDUCTION: govtfund_5 — keep 5 of 10 items (items 3, 4, 5, 6, 7 in bold; others highlighted for removal)",
         bold=True, size=10, color=RGBColor(0xCC, 0x00, 0x00))
add_qvar(doc, "govtfund_5",
         "How likely is it that these things happen when the government funds academic research?")
for num, text in g5_items:
    cut = num in G5_CUT
    add_bullet(doc, text, highlight=cut, strikethrough=cut)
add_scale_note(doc, "5-point scale: Extremely unlikely / Somewhat unlikely / Neither likely nor unlikely / Somewhat likely / Extremely likely")

doc.add_paragraph()
add_qvar(doc, "gotvfund_6", "Rate your agreement with the following statements:")
for item in [
    "If the government funds research, there should be rules to ensure that the research is done well",
    "University oversight committees protect against unethical research",
    "Many everyday tools are the result of academic research",
    "Research at universities benefits everyone's day to day lives in ways that are difficult to see",
    "University researchers are motivated primarily by profit",
    "Private companies conduct research primarily to benefit the public",
    "Government-funded research is more likely to be objective than research funded by private companies",
    "Research funded by private companies is more likely to be biased toward results that benefit that company",
    "Results of research by private companies will become publicly accessible",
]:
    add_bullet(doc, item)
add_scale_note(doc, "5-point scale: Strongly disagree / Somewhat disagree / Neither agree nor disagree / Somewhat agree / Strongly agree")

doc.add_paragraph()
add_qvar(doc, "govtfund_7", "What are the priorities of government-funded research at colleges and universities? Rate how highly prioritized each is:")
for item in ["Making a profit", "Advance knowledge and benefit society", "Develop products that can be sold", "Serve the interests of the government"]:
    add_bullet(doc, item)
add_scale_note(doc, "5-point scale: Not a priority / Low priority / Somewhat priority / Moderate priority / High priority")

doc.add_paragraph()
add_qvar(doc, "govtfund_8", "Rate your agreement with the following statements: Investing tax dollars in academic research...")
for item in ["...boosts the economy", "...improves health and saves lives", "...drives industry investment", "...makes America a leader", "...protects national security", "...leads to scientific breakthroughs", "...benefits people like you"]:
    add_bullet(doc, item)
add_scale_note(doc, "5-point scale: Strongly disagree / Somewhat disagree / Neither agree nor disagree / Somewhat agree / Strongly agree")

doc.add_paragraph()
add_qvar(doc, "govtfund_9", "Would you approve or disapprove if the federal government took the following actions?")
for item in ["Cutting funding for health research", "Cutting funding for education", "Cutting funding for science", "Cutting funding for the arts", "Cutting funding for the humanities"]:
    add_bullet(doc, item)
add_scale_note(doc, "5-point scale: Strongly disapprove / Somewhat disapprove / Neither approve nor disapprove / Somewhat approve / Strongly approve")

doc.add_paragraph()
add_qvar(doc, "govfund_10", "Do you think your local area would be positively or negatively impacted by cuts to funding for scientific research at universities?")
for opt in ["Strong negative impact", "Somewhat negative impact", "Neither positive nor negative impact", "Somewhat positive impact", "Strong positive impact"]:
    add_bullet(doc, opt)

# ════════════════════════════════════════════════════════════════════════════
#  BLOCK 4: ATTENTION CHECK 1
# ════════════════════════════════════════════════════════════════════════════
section_break(doc, "Block: Attention Check 1")
add_qvar(doc, "attention_1", 'Please select "Somewhat agree" below if you are paying attention.')
for opt in ["Disagree", "Somewhat disagree", "Neither agree nor disagree", "Somewhat agree", "Agree"]:
    add_bullet(doc, opt)

# ════════════════════════════════════════════════════════════════════════════
#  BLOCK 5: BEHAVIOR INTENTIONS
# ════════════════════════════════════════════════════════════════════════════
section_break(doc, "Block: Behavior Intentions")
add_para(doc, "Intro: Please rate whether or not you intend to take the following actions:", italic=True, size=9, color=GRAY)
add_qvar(doc, "intent (×4 items)", "[Four intention items — text varies by Qualtrics loop/merge field]")
add_scale_note(doc, "7-point scale: Definitely will not / Probably will not / Maybe will not / Unsure / Maybe will / Probably will / Definitely will")

# ════════════════════════════════════════════════════════════════════════════
#  BLOCK 6: PAST ENGAGEMENT
# ════════════════════════════════════════════════════════════════════════════
section_break(doc, "Block: Past Engagement")
add_para(doc, "Intro: The next set of questions will ask about how you have engaged with science and research in the past.", italic=True, size=9, color=GRAY)
doc.add_paragraph()
add_qvar(doc, "past_engage", "Over the past 12 months, how often have you done the following:")
for item in [
    "Had conversations with friends, family, or co-workers about science-related topics (e.g., climate change, vaccination, nutrition)",
    "Shared or commented on social media posts about scientific issues",
    "Attended public rallies or protests related to scientific issues (e.g., 'Stand Up for Science', 'March for Science', 'Fridays for Future')",
    "Shared information online (including on social media) about the cuts to federal funding for scientific research",
    "Talked to other people (like your friends and family) about the cuts to federal funding for scientific research",
    "Contacted your representatives in Congress to ask them to oppose cuts to federal funding for research at universities",
    "Contacted your representatives in Congress to ask them to support federal funding for research at universities",
]:
    add_bullet(doc, item)
add_scale_note(doc, "6-point scale: Never / Once or twice a year / Once or twice a month / Once or twice a week / Almost every day / Once or more per day")

# ════════════════════════════════════════════════════════════════════════════
#  BLOCK 7: ATTENTION CHECK 2
# ════════════════════════════════════════════════════════════════════════════
section_break(doc, "Block: Attention Check 2")
add_qvar(doc, "attention_2", 'Please select "Disagree" below if you are paying attention.')
for opt in ["Disagree", "Somewhat disagree", "Neither agree nor disagree", "Somewhat agree", "Agree"]:
    add_bullet(doc, opt)

# ════════════════════════════════════════════════════════════════════════════
#  BLOCK 8: HIGHER ED
# ════════════════════════════════════════════════════════════════════════════
section_break(doc, "Block: Higher Education")

add_qvar(doc, "highered_1", "Overall, how much do you approve or disapprove of:")
for item in ["Private colleges and universities in the United States", "Your state's public university system", "Your local community college(s)"]:
    add_bullet(doc, item)
add_scale_note(doc, "5-point scale: Strongly disapprove / Somewhat disapprove / Neutral / Somewhat approve / Strongly approve")

doc.add_paragraph()
add_qvar(doc, "highered_2", "How much do you trust each type of institution to do what is right?")
for item in ["Private colleges and universities in the United States", "Your state's public university system", "Your local community college(s)"]:
    add_bullet(doc, item)
add_scale_note(doc, "5-point scale: None at all / A little / A moderate amount / A lot / A great deal")

doc.add_paragraph()
add_qvar(doc, "highered_3", "How much do each of the following contribute to scientific achievements in the US?")
for item in ["Private colleges and universities", "Private companies (pharmaceutical companies, technology companies, etc.)", "Federal government agencies", "Charitable foundations", "State colleges and universities", "Community colleges"]:
    add_bullet(doc, item)
add_scale_note(doc, "6-point scale: None at all / A little / A moderate amount / A lot / A great deal / Not sure")

doc.add_paragraph()
add_qvar(doc, "highered_4", "What proportion of each type of institution do you think is nonprofit?")
for item in ["Private colleges and universities", "Public state colleges and universities", "Community colleges", "Private companies (e.g. pharmaceutical and tech companies)", "Federal government agencies", "Charitable foundations"]:
    add_bullet(doc, item)
add_scale_note(doc, "6-point scale: None / Less than half / About half / Most / All / Don't know")

# ════════════════════════════════════════════════════════════════════════════
#  BLOCK 9: UNIVERSITIES
# ════════════════════════════════════════════════════════════════════════════
section_break(doc, "Block: Universities")
add_para(doc, "Intro: In the next few questions, we are asking for your opinions specifically about 4-year colleges and universities.", italic=True, size=9, color=GRAY)
doc.add_paragraph()

add_qvar(doc, "univ_2", "To what extent do you agree or disagree that colleges and universities contribute to each of the following?")
for item in ["Greater innovation, including new scientific, medical, and technological discoveries", "Better jobs, promotions, and career advancements", "Higher household incomes", "A more knowledgeable population", "A competitive advantage for the US over other countries", "More entrepreneurship and business creation", "Increased compassion for, and tolerance of, others", "Overall economic growth", "Better relationships with other countries", "Benefits to surrounding communities"]:
    add_bullet(doc, item)
add_scale_note(doc, "5-point scale: Strongly disagree / Somewhat disagree / Neither agree nor disagree / Somewhat agree / Strongly agree")

doc.add_paragraph()
add_qvar(doc, "univ_4", "Do you agree or disagree that colleges and universities:")
for item in ["Discover new knowledge through original research", "Apply existing knowledge and connect ideas across contexts and disciplines", "Apply knowledge-based solutions to society's needs"]:
    add_bullet(doc, item)
add_scale_note(doc, "5-point scale: Strongly disagree / Somewhat disagree / Neither agree nor disagree / Somewhat agree / Strongly agree")

doc.add_paragraph()
add_qvar(doc, "univ_5", "How much do each of the following people benefit from colleges and universities?")
for item in ["You personally", "People you care about", "Local communities", "Students who attend", "Students' families", "Employers and businesses", "Society as a whole"]:
    add_bullet(doc, item)
add_scale_note(doc, "5-point scale: None at all / A little / A moderate amount / A lot / A great deal")

# ════════════════════════════════════════════════════════════════════════════
#  BLOCK 10: ATTENTION CHECK 3
# ════════════════════════════════════════════════════════════════════════════
section_break(doc, "Block: Attention Check 3")
add_qvar(doc, "attention_3", 'Please select "Somewhat disagree" below if you are paying attention.')
for opt in ["Disagree", "Somewhat disagree", "Neither agree nor disagree", "Somewhat agree", "Agree"]:
    add_bullet(doc, opt)

# ════════════════════════════════════════════════════════════════════════════
#  BLOCK 11: CLIMATE
# ════════════════════════════════════════════════════════════════════════════
section_break(doc, "Block: Climate")
add_para(doc, "Intro: In the next set of questions, we will ask you about climate change, the long-term shift in global weather patterns and average temperatures.", italic=True, size=9, color=GRAY)
doc.add_paragraph()

add_qvar(doc, "climate_1", "Do you think global warming is happening?")
for opt in ["No", "Yes", "Don't know"]:
    add_bullet(doc, opt)

doc.add_paragraph()
add_qvar(doc, "climate_2", "Assuming global warming is happening, do you think it is...")
for opt in ["Caused mostly by humans", "Caused mostly by natural changes in the environment", "None of the above, because global warming isn't happening", "Other"]:
    add_bullet(doc, opt)

doc.add_paragraph()
add_qvar(doc, "climate_3", "To what extent do you trust scientists who work on climate change?")
for opt in ["Not at all", "Slightly", "Moderately", "Very strongly"]:
    add_bullet(doc, opt)

doc.add_paragraph()
add_qvar(doc, "climate_4", "How important do you think it is for colleges and universities to do research on the following climate-related topics?")
for item in [
    "Climate science and earth systems (such as tracking greenhouse gases, studying rising sea levels, and understanding past climates)",
    "Climate mitigation approaches ('fixes' such as renewable energy, carbon capture, and sustainable agriculture)",
    "Climate adaptation and resilience ('adjustments' such as urban infrastructure, carbon-smart crops, and ecosystem management)",
    "Climate policy, economics, and human impact (such as environmental justice, climate finance, and psychology and communication)",
]:
    add_bullet(doc, item)
add_scale_note(doc, "5-point scale: Not at all important / Slightly important / Moderately important / Very important / Extremely important")

# ════════════════════════════════════════════════════════════════════════════
#  BLOCK 12: DEMOGRAPHICS
# ════════════════════════════════════════════════════════════════════════════
section_break(doc, "Block: Demographics")

add_qvar(doc, "zipcode", "What is your zip code? [Open text]")
doc.add_paragraph()

add_qvar(doc, "gender", "Which best describes your gender identity?")
for opt in ["Man", "Woman", "Non-binary", "Genderqueer", "Agender", "Gender fluid", "I prefer to self-describe", "Prefer not to say"]:
    add_bullet(doc, opt)

doc.add_paragraph()
add_qvar(doc, "LGBTQ", "Do you identify as part of the LGBTQIA+ community?")
for opt in ["Yes", "No", "Prefer not to say"]:
    add_bullet(doc, opt)

doc.add_paragraph()
add_qvar(doc, "race", "What race/ethnicity or races/ethnicities do you consider yourself to be? (Select all that apply)")
for opt in [
    "White (e.g. German, Irish, English, Italian, Polish, French)",
    "Black or African American (e.g. Jamaican, Haitian, Nigerian, Ethiopian, Somalian)",
    "East Asian (e.g. Chinese, Japanese, Korean, Taiwanese)",
    "American Indian or Alaska Native",
    "Native Hawaiian or Other Pacific Islander",
    "South Asian (e.g. Asian Indian, Pakistani, Bangladeshi, Nepalese, Sri Lankan)",
    "Middle Eastern or North African (e.g. Lebanese, Iranian, Egyptian, Syrian, Moroccan, Algerian)",
    "Prefer not to say",
    "My racial/ethnic identity is not listed here",
]:
    add_bullet(doc, opt)

doc.add_paragraph()
add_qvar(doc, "ses_subjective", "Think of this ladder below as representing where people stand in the United States. [Subjective SES ladder, 1–10]")
add_scale_note(doc, "10-point slider: 1 (bottom) to 10 (top)")

doc.add_paragraph()
add_qvar(doc, "univ_affiliate", "Are you currently affiliated with a college or university? (Select all that apply)")
for opt in ["Current student", "Past student (alumni)", "Staff", "Faculty", "No affiliation", "Other affiliation"]:
    add_bullet(doc, opt)

doc.add_paragraph()
add_qvar(doc, "ses_degree", "What is the highest degree or level of school you have completed?")
for opt in [
    "No formal education", "Did not complete high school",
    "High school graduate (diploma)", "High school graduate (GED)",
    "Some college (1-4 years, no degree)",
    "Associate's degree (including occupational or academic degrees)",
    "Bachelor's degree (BA, BS, etc.)", "Master's degree (MA, MS, MEng, MSW, etc.)",
    "Professional school degree (MD, DDC, JD, etc.)", "Doctoral degree (PhD, EdD, etc.)",
]:
    add_bullet(doc, opt)

doc.add_paragraph()
add_qvar(doc, "ses_income_household", "Please indicate your approximate total household income before taxes in the past 12 months.")
for opt in ["Less than $5,000", "$5,000–$11,999", "$12,000–$15,999", "$16,000–$24,999",
            "$25,000–$34,999", "$35,000–$49,999", "$50,000–$74,999", "$75,000–$99,999",
            "$100,000–$149,999", "$150,000 or greater", "Prefer not to say", "I don't know"]:
    add_bullet(doc, opt)

doc.add_paragraph()
add_qvar(doc, "politics_ideology", "Here is a seven-point scale on which the political views that people might hold are arranged from extremely liberal to extremely conservative.")
for opt in ["Extremely liberal", "Liberal", "Slightly liberal", "Moderate / middle of the road", "Slightly conservative", "Conservative", "Extremely conservative"]:
    add_bullet(doc, opt)

doc.add_paragraph()
add_qvar(doc, "urbanicity", "Which of the following best describes the area where you live?")
for opt in ["Urban (city center, high population density)", "Suburban (residential area near a city)", "Rural (countryside, low population density)"]:
    add_bullet(doc, opt)

# ════════════════════════════════════════════════════════════════════════════
#  BLOCK 13: END
# ════════════════════════════════════════════════════════════════════════════
section_break(doc, "Block: End")
add_qvar(doc, "end_feedback", "Please type below any comments or suggestions you might have for the research team. [Open text]")
add_qvar(doc, "end_next", "When you are done, click 'Next' at the bottom right to be redirected to Prolific and register your completion.")

# ── save ────────────────────────────────────────────────────────────────────
out = "UnivNews_Step1_Survey_Readable.docx"
doc.save(out)
print(f"Saved: {out}")
